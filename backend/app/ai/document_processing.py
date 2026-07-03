from __future__ import annotations

import hashlib
import io
import logging
import re
import time
import unicodedata
import zipfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from langchain_core.documents import Document


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".md", ".txt"}
logger = logging.getLogger("uvicorn.error")

STRICT_OOXML_NAMESPACE_REPLACEMENTS = (
    (
        b"http://purl.oclc.org/ooxml/officeDocument/relationships",
        b"http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    ),
    (
        b"http://purl.oclc.org/ooxml/wordprocessingml/main",
        b"http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    ),
    (
        b"http://purl.oclc.org/ooxml/drawingml/wordprocessingDrawing",
        b"http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
    ),
    (
        b"http://purl.oclc.org/ooxml/drawingml/main",
        b"http://schemas.openxmlformats.org/drawingml/2006/main",
    ),
    (
        b"http://purl.oclc.org/ooxml/officeDocument/math",
        b"http://schemas.openxmlformats.org/officeDocument/2006/math",
    ),
    (
        b"http://purl.oclc.org/ooxml/officeDocument/customProperties",
        b"http://schemas.openxmlformats.org/officeDocument/2006/custom-properties",
    ),
    (
        b"http://purl.oclc.org/ooxml/officeDocument/extendedProperties",
        b"http://schemas.openxmlformats.org/officeDocument/2006/extended-properties",
    ),
    (
        b"http://purl.oclc.org/ooxml/officeDocument/docPropsVTypes",
        b"http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes",
    ),
)


class DocumentProcessingError(ValueError):
    """Raised when an uploaded document cannot be prepared for indexing."""


@dataclass(frozen=True)
class ProcessedDocument:
    documents: list[Document]
    warnings: list[str] = field(default_factory=list)


def validate_document_extension(filename: str | Path) -> str:
    """校验文档扩展名并返回标准化的小写后缀。"""

    suffix = Path(filename).suffix.lower()
    if suffix in SUPPORTED_EXTENSIONS:
        return suffix

    file_type = suffix or "无扩展名"
    supported = "、".join(sorted(SUPPORTED_EXTENSIONS))
    suggestion = "；旧版 .doc 文件请先转换为 .docx" if suffix == ".doc" else ""
    raise DocumentProcessingError(
        f"暂不支持 {file_type} 文件，当前支持格式：{supported}{suggestion}"
    )


def clean_text(text: str) -> str:
    """Normalize extracted text without destroying paragraph boundaries."""

    normalized = unicodedata.normalize("NFKC", text)
    normalized = normalized.replace("\r\n", "\n").replace("\r", "\n")
    normalized = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", normalized)
    normalized = re.sub(r"[ \t]+", " ", normalized)
    normalized = re.sub(r" *\n *", "\n", normalized)
    normalized = re.sub(r"\n{3,}", "\n\n", normalized)

    return normalized.strip()


def load_upload_documents(
    file_path: Path,
    *,
    original_filename: str | None = None,
    content_type: str | None = None,
) -> ProcessedDocument:
    """Extract, clean, and wrap an uploaded file as LangChain documents."""

    path = Path(file_path)
    suffix = validate_document_extension(path)
    parser = {
        ".pdf": "pypdf",
        ".docx": "python-docx",
        ".md": "plain-text",
        ".txt": "plain-text",
    }[suffix]
    logger.info(
        "[document] Parser selected file=%r type=%s content_type=%s parser=%s size=%s bytes",
        original_filename or path.name,
        suffix,
        content_type or "unknown",
        parser,
        path.stat().st_size,
    )

    if suffix == ".pdf":
        return _load_pdf_documents(
            path,
            original_filename=original_filename,
            content_type=content_type,
        )

    if suffix == ".docx":
        return _load_docx_documents(
            path,
            original_filename=original_filename,
            content_type=content_type,
        )

    return _load_text_document(
        path,
        original_filename=original_filename,
        content_type=content_type,
    )


def _load_pdf_documents(
    path: Path,
    *,
    original_filename: str | None,
    content_type: str | None,
) -> ProcessedDocument:
    """Extract text-layer PDF pages without loading OCR/inference dependencies."""

    display_name = original_filename or path.name
    import_started_at = time.perf_counter()
    logger.info("[document] PDF parser import started file=%r parser=pypdf", display_name)
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        logger.exception(
            "[document] PDF parser import failed file=%r elapsed=%.2fs error=%s",
            display_name,
            time.perf_counter() - import_started_at,
            exc,
        )
        raise DocumentProcessingError("缺少 pypdf 依赖，请先安装 AI 可选依赖") from exc
    logger.info(
        "[document] PDF parser import completed file=%r elapsed=%.2fs",
        display_name,
        time.perf_counter() - import_started_at,
    )

    started_at = time.perf_counter()
    logger.info(
        "[document] Parsing started file=%r type=.pdf parser=pypdf size=%s bytes",
        display_name,
        path.stat().st_size,
    )
    try:
        reader = PdfReader(path)
    except Exception as exc:
        logger.exception(
            "[document] Parsing failed file=%r elapsed=%.2fs error=%s",
            display_name,
            time.perf_counter() - started_at,
            exc,
        )
        raise DocumentProcessingError(f"文档解析失败：{exc}") from exc

    documents: list[Document] = []
    warnings: list[str] = []
    file_hash = _file_sha256(path)
    empty_page_count = 0
    total_characters = 0

    for index, page in enumerate(reader.pages):
        text = clean_text(page.extract_text() or "")
        if not text:
            empty_page_count += 1
            continue

        total_characters += len(text)

        metadata = _clean_metadata({
            "source": str(path),
            "original_filename": display_name,
            "content_type": content_type,
            "file_ext": ".pdf",
            "element_index": index,
            "element_type": "Page",
            "page_number": index + 1,
            "file_sha256": file_hash,
            "parser": "pypdf",
        })

        documents.append(Document(page_content=text, metadata=metadata))

    if not documents:
        warnings.append("未提取到可用文本，当前版本暂未启用 OCR")
        logger.warning(
            "[document] Parsing produced no text file=%r pages=%s elapsed=%.2fs",
            display_name,
            len(reader.pages),
            time.perf_counter() - started_at,
        )
        raise DocumentProcessingError(warnings[0])

    logger.info(
        "[document] Parsing completed file=%r parser=pypdf pages_total=%s "
        "pages_with_text=%s pages_empty=%s characters=%s sha256=%s elapsed=%.2fs",
        display_name,
        len(reader.pages),
        len(documents),
        empty_page_count,
        total_characters,
        file_hash[:12],
        time.perf_counter() - started_at,
    )
    return ProcessedDocument(documents=documents, warnings=warnings)


def _load_docx_documents(
    path: Path,
    *,
    original_filename: str | None,
    content_type: str | None,
) -> ProcessedDocument:
    """使用 python-docx 直接提取段落和表格，避免自动解析器的重型导入。"""

    display_name = original_filename or path.name
    import_started_at = time.perf_counter()
    logger.info(
        "[document] DOCX parser import started file=%r parser=python-docx",
        display_name,
    )
    try:
        from docx import Document as load_docx
    except ImportError as exc:
        logger.exception(
            "[document] DOCX parser import failed file=%r elapsed=%.2fs error=%s",
            display_name,
            time.perf_counter() - import_started_at,
            exc,
        )
        raise DocumentProcessingError("缺少 python-docx 依赖，请先安装 AI 可选依赖") from exc
    logger.info(
        "[document] DOCX parser import completed file=%r elapsed=%.2fs",
        display_name,
        time.perf_counter() - import_started_at,
    )

    started_at = time.perf_counter()
    logger.info(
        "[document] Parsing started file=%r type=.docx parser=python-docx size=%s bytes",
        display_name,
        path.stat().st_size,
    )
    try:
        docx_source, strict_ooxml = _prepare_docx_source(path)
        if strict_ooxml:
            logger.info(
                "[document] Strict OOXML detected file=%r action=normalize_namespaces",
                display_name,
            )
        word_document = load_docx(docx_source)
    except Exception as exc:
        logger.exception(
            "[document] Parsing failed file=%r parser=python-docx elapsed=%.2fs error=%s",
            display_name,
            time.perf_counter() - started_at,
            exc,
        )
        raise DocumentProcessingError(f"文档解析失败：{exc}") from exc

    documents: list[Document] = []
    warnings: list[str] = []
    file_hash = _file_sha256(path)
    total_characters = 0
    paragraph_count = 0
    table_cell_count = 0

    def append_text(text: str, element_type: str) -> None:
        nonlocal total_characters, paragraph_count, table_cell_count
        cleaned = clean_text(text)
        if not cleaned:
            return

        element_index = len(documents)
        total_characters += len(cleaned)
        if element_type == "Paragraph":
            paragraph_count += 1
        else:
            table_cell_count += 1

        documents.append(
            Document(
                page_content=cleaned,
                metadata=_clean_metadata({
                    "source": str(path),
                    "original_filename": display_name,
                    "content_type": content_type,
                    "file_ext": ".docx",
                    "element_index": element_index,
                    "element_type": element_type,
                    "file_sha256": file_hash,
                    "parser": "python-docx",
                }),
            )
        )

    for paragraph in word_document.paragraphs:
        append_text(paragraph.text, "Paragraph")

    seen_cells: set[int] = set()
    for table in word_document.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_id = id(cell._tc)
                if cell_id in seen_cells:
                    continue
                seen_cells.add(cell_id)
                append_text(cell.text, "TableCell")

    if not documents:
        warnings.append("未提取到可用文本")
        logger.warning(
            "[document] Parsing produced no text file=%r parser=python-docx elapsed=%.2fs",
            display_name,
            time.perf_counter() - started_at,
        )
        raise DocumentProcessingError(warnings[0])

    logger.info(
        "[document] Parsing completed file=%r parser=python-docx paragraphs=%s "
        "table_cells=%s documents=%s characters=%s sha256=%s elapsed=%.2fs",
        display_name,
        paragraph_count,
        table_cell_count,
        len(documents),
        total_characters,
        file_hash[:12],
        time.perf_counter() - started_at,
    )
    return ProcessedDocument(documents=documents, warnings=warnings)


def _load_text_document(
    path: Path,
    *,
    original_filename: str | None,
    content_type: str | None,
) -> ProcessedDocument:
    """直接读取 Markdown 和纯文本文件，并兼容常见中文编码。"""

    display_name = original_filename or path.name
    started_at = time.perf_counter()
    logger.info(
        "[document] Parsing started file=%r type=%s parser=plain-text size=%s bytes",
        display_name,
        path.suffix.lower(),
        path.stat().st_size,
    )

    text = None
    selected_encoding = None
    for encoding in ("utf-8-sig", "gb18030"):
        try:
            text = path.read_text(encoding=encoding)
            selected_encoding = encoding
            break
        except UnicodeDecodeError:
            continue
        except OSError as exc:
            raise DocumentProcessingError(f"文档读取失败：{exc}") from exc

    if text is None or selected_encoding is None:
        raise DocumentProcessingError("文档编码不受支持，请使用 UTF-8 或 GB18030 编码")

    cleaned = clean_text(text)
    if not cleaned:
        raise DocumentProcessingError("未提取到可用文本")

    file_hash = _file_sha256(path)
    document = Document(
        page_content=cleaned,
        metadata=_clean_metadata({
            "source": str(path),
            "original_filename": display_name,
            "content_type": content_type,
            "file_ext": path.suffix.lower(),
            "element_index": 0,
            "element_type": "Text",
            "file_sha256": file_hash,
            "parser": "plain-text",
            "encoding": selected_encoding,
        }),
    )
    logger.info(
        "[document] Parsing completed file=%r parser=plain-text encoding=%s "
        "characters=%s sha256=%s elapsed=%.2fs",
        display_name,
        selected_encoding,
        len(cleaned),
        file_hash[:12],
        time.perf_counter() - started_at,
    )
    return ProcessedDocument(documents=[document])


def _prepare_docx_source(path: Path) -> tuple[str | io.BytesIO, bool]:
    """将 Strict OOXML 命名空间转换为 python-docx 支持的 Transitional OOXML。"""

    try:
        with zipfile.ZipFile(path) as source_zip:
            root_relationships = source_zip.read("_rels/.rels")
            is_strict = b"http://purl.oclc.org/ooxml/" in root_relationships
            if not is_strict:
                return str(path), False

            converted_stream = io.BytesIO()
            with zipfile.ZipFile(converted_stream, "w") as target_zip:
                for member in source_zip.infolist():
                    content = source_zip.read(member.filename)
                    if member.filename.endswith((".xml", ".rels")):
                        for strict_namespace, transitional_namespace in (
                            STRICT_OOXML_NAMESPACE_REPLACEMENTS
                        ):
                            content = content.replace(
                                strict_namespace,
                                transitional_namespace,
                            )
                    target_zip.writestr(member, content)

            converted_stream.seek(0)
            return converted_stream, True
    except (KeyError, zipfile.BadZipFile, OSError) as exc:
        raise DocumentProcessingError(f"DOCX 文件结构无效或已损坏：{exc}") from exc


def _file_sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as file:
        for chunk in iter(lambda: file.read(1024 * 1024), b""):
            digest.update(chunk)

    return digest.hexdigest()


def _clean_metadata(metadata: dict[str, Any]) -> dict[str, Any]:
    return {
        key: value
        for key, value in metadata.items()
        if value is not None and isinstance(value, str | int | float | bool)
    }
