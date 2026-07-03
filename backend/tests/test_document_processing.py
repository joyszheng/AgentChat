import io
import sys
import zipfile
from types import ModuleType

import pytest

from app.ai.document_processing import (
    DocumentProcessingError,
    clean_text,
    load_upload_documents,
)


def test_clean_text_normalizes_whitespace():
    text = "  第一段\r\n\r\n\r\n  第二段\t内容\x00  "

    assert clean_text(text) == "第一段\n\n第二段 内容"


def test_load_upload_documents_rejects_unsupported_extension(tmp_path):
    file_path = tmp_path / "data.xlsx"
    file_path.write_text("name,value", encoding="utf-8")

    with pytest.raises(DocumentProcessingError, match="暂不支持"):
        load_upload_documents(file_path)


def test_load_upload_documents_uses_pypdf_for_pdf(monkeypatch, tmp_path):
    file_path = tmp_path / "guide.pdf"
    file_path.write_text("# 标题\n\n正文", encoding="utf-8")

    calls = []

    class FakePage:
        def extract_text(self):
            calls.append("extract_text")
            return "  文档正文\r\n\r\n"

    class FakePdfReader:
        def __init__(self, path):
            calls.append(path)
            self.pages = [FakePage()]

    pypdf_module = ModuleType("pypdf")
    pypdf_module.PdfReader = FakePdfReader

    monkeypatch.setitem(sys.modules, "pypdf", pypdf_module)

    processed = load_upload_documents(
        file_path,
        original_filename="guide.pdf",
        content_type="application/pdf",
    )

    assert calls == [file_path, "extract_text"]
    assert len(processed.documents) == 1
    assert processed.documents[0].page_content == "文档正文"
    assert processed.documents[0].metadata["original_filename"] == "guide.pdf"
    assert processed.documents[0].metadata["page_number"] == 1
    assert processed.documents[0].metadata["parser"] == "pypdf"


def test_load_upload_documents_uses_python_docx_directly(tmp_path):
    from docx import Document as WordDocument

    file_path = tmp_path / "guide.docx"
    word_document = WordDocument()
    word_document.add_paragraph("第一段正文")
    table = word_document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "表格内容"
    word_document.save(file_path)

    processed = load_upload_documents(
        file_path,
        original_filename="guide.docx",
        content_type=(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ),
    )

    assert [document.page_content for document in processed.documents] == [
        "第一段正文",
        "表格内容",
    ]
    assert {document.metadata["parser"] for document in processed.documents} == {
        "python-docx"
    }


def test_load_upload_documents_reads_gb18030_text(tmp_path):
    file_path = tmp_path / "guide.txt"
    file_path.write_bytes("中文文本".encode("gb18030"))

    processed = load_upload_documents(file_path)

    assert processed.documents[0].page_content == "中文文本"
    assert processed.documents[0].metadata["parser"] == "plain-text"
    assert processed.documents[0].metadata["encoding"] == "gb18030"


def test_load_upload_documents_supports_strict_ooxml_docx(tmp_path):
    from docx import Document as WordDocument

    transitional_path = tmp_path / "transitional.docx"
    strict_path = tmp_path / "strict.docx"
    word_document = WordDocument()
    word_document.add_paragraph("Strict OOXML 正文")
    word_document.save(transitional_path)

    converted_stream = io.BytesIO()
    with zipfile.ZipFile(transitional_path) as source_zip:
        with zipfile.ZipFile(converted_stream, "w") as target_zip:
            for member in source_zip.infolist():
                content = source_zip.read(member.filename)
                if member.filename.endswith((".xml", ".rels")):
                    content = content.replace(
                        b"http://schemas.openxmlformats.org/officeDocument/2006/relationships",
                        b"http://purl.oclc.org/ooxml/officeDocument/relationships",
                    ).replace(
                        b"http://schemas.openxmlformats.org/wordprocessingml/2006/main",
                        b"http://purl.oclc.org/ooxml/wordprocessingml/main",
                    )
                target_zip.writestr(member, content)
    strict_path.write_bytes(converted_stream.getvalue())

    processed = load_upload_documents(strict_path)

    assert processed.documents[0].page_content == "Strict OOXML 正文"
    assert processed.documents[0].metadata["parser"] == "python-docx"
